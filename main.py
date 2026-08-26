from cliente import Cliente
from prodotto import Prodotto
from ordine import Ordine

import os
import tkinter as tk
from tkinter import messagebox, ttk, PhotoImage
import mysql.connector
import threading
import json
import csv
from docx import Document
from docx.shared import Inches
from pathlib import Path
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
from sklearn.cluster import KMeans
import random
import itertools
from sklearn.linear_model import LinearRegression
from datetime import date, datetime
import time
import schedule
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import bcrypt
from dotenv import load_dotenv

load_dotenv()

conn = mysql.connector.connect(
    host = os.getenv("DB_HOST"),
    user = os.getenv("DB_USER"),
    password = os.getenv("DB_PASS"),
    database = os.getenv("DB_NAME"),
    port = int(os.getenv("DB_PORT"))
)
cursor = conn.cursor(dictionary=True)

FILE_PATH = Path(__file__).resolve().parent
WINDOW_SIZE = "1920x1080"
BG="#F5F5F5" 
FONT="Cambria"
ENCODING= "utf-8"
FAVICON= FILE_PATH / "favicon.ico"
PADDING_X = 2
PADDING_Y= 2

root = tk.Tk()
root.title("Gestionale E-commerce")
root.geometry(WINDOW_SIZE)
root.config(bg = BG) 
root.iconbitmap(FAVICON)

df_clienti= None
df_prodotti= None
df_ordini= None
df_completo= None
canvas = None

def addWidget(element, row, column, rowspan, columnspan):   
    element.grid(row=row, column=column, rowspan=rowspan, columnspan=columnspan, padx=PADDING_X, pady=PADDING_Y, sticky="nsew")

def manda_email(corpo):
    mittente= os.getenv("SMTP_USER")
    destinatario= os.getenv("SMTP_DEST")
    oggetto= "Backup dati"
    password= os.getenv("SMTP_PASS")
    messaggio= MIMEMultipart()
    messaggio['From'] = mittente
    messaggio['To'] = destinatario
    messaggio['Subject'] = oggetto
    messaggio.attach(MIMEText(corpo, 'plain'))
    server = smtplib.SMTP('smtp.gmail.com', 587)
    server.starttls()
    server.login(mittente, password)
    server.sendmail(mittente, destinatario, messaggio.as_string())  
    server.quit()

def esegui_funz(funz, nome):
    try:
        funz()
        manda_email(f"{nome} eseguito correttamente")
    except Exception as e:
        manda_email(f"Errore durante l'esecuzione del backup: {nome}")

def backup():
    schedule.every().day.at("17:30").do(lambda: esegui_funz(esportare_dati, "Esportazione dati"))
    schedule.every().day.at("18:00").do(lambda: esegui_funz(creazione_json, "Creazione JSON"))
    while True:
        schedule.run_pending()
        time.sleep(60)

def crea_df_clienti():
    global df_clienti
    cursor.execute("SELECT * FROM clienti")
    risultati = cursor.fetchall()
    df_clienti = pd.DataFrame(risultati, columns=['id_cliente', 'nome', 'cognome', 'email', 'password', 'indirizzo_consegna'])
    return df_clienti

def crea_df_prodotti():
    global df_prodotti
    cursor.execute("SELECT * FROM prodotti")
    risultati = cursor.fetchall()
    df_prodotti = pd.DataFrame(risultati, columns=['id_prodotto', 'tipo', 'categoria', 'descrizione', 'prezzo_unitario', 'quantita_magazzino'])
    df_prodotti['prezzo_unitario']= pd.to_numeric(df_prodotti['prezzo_unitario'], errors="coerce") 
    df_prodotti['quantita_magazzino']= pd.to_numeric(df_prodotti['quantita_magazzino'], errors="coerce") 
    return df_prodotti

def crea_df_ordini():
    global df_ordini
    cursor.execute("SELECT * FROM ordini")
    risultati = cursor.fetchall()
    df_ordini = pd.DataFrame(risultati, columns=['id_ordine', 'id_prodotto', 'id_cliente', 'quantita_venduta', 'importo', 'data_ordine'])
    df_ordini['quantita_venduta']= pd.to_numeric(df_ordini['quantita_venduta'], errors="coerce") 
    df_ordini['importo']= pd.to_numeric(df_ordini['importo'], errors="coerce") 
    df_ordini['data_ordine']= pd.to_datetime(df_ordini['data_ordine'], errors="coerce")
    return df_ordini

def esportare_dati():
    crea_df()
    directory= os.path.join(FILE_PATH / "report" / str(date.today())) 
    os.makedirs(directory, exist_ok=True)
    df_clienti.to_csv(os.path.join(directory, "dati clienti.csv"), index=False)
    df_prodotti.to_csv(os.path.join(directory,"dati prodotti.csv"), index=False)
    df_ordini.to_csv(os.path.join(directory,"dati ordini.csv"), index=False)

def unione_dati():
    global df_completo
    if df_clienti is None or df_prodotti is None or df_ordini is None:
        messagebox.showerror("Errore", "Dati non trovati")
        return
    df_merge = df_ordini.merge(df_prodotti, on="id_prodotto", how="left")
    df_completo = pd.merge(df_merge, df_clienti, on="id_cliente", how="left")
    return df_completo

lock= threading.Lock()
def crea_df():
    with lock:
        crea_df_clienti()
        crea_df_prodotti()
        crea_df_ordini()
        unione_dati()
crea_df()

def creazione_json():
    crea_df()
    directory= os.path.join(FILE_PATH / "file_json") 
    os.makedirs(directory, exist_ok=True)
    df_clienti.to_json(os.path.join(directory,"elenco_clienti.json"), indent=4, ensure_ascii= False)
    df_prodotti.to_json(os.path.join(directory,"elenco_prodotti.json"), indent=4, ensure_ascii= False)
    df_ordini.to_json(os.path.join(directory,"elenco_ordini.json"), indent=4, ensure_ascii= False)

def valida_email(email):
    if "@" not in email or "." not in email:
        messagebox.showerror("Errore", "Email non valida")
        return False
    cursor.execute("SELECT * FROM clienti WHERE email = %s", (email,))
    if cursor.fetchone():
        messagebox.showerror("Errore", "Email già registrata")
        return False
    return email

def valida_password(password):
    if len(password) < 8:
        messagebox.showerror("Errore", "La password deve contenere almeno 8 caratteri")
        return False
    if not any(c.isdigit() for c in password):
        messagebox.showerror("Errore", "La password deve contenere almeno un numero")
        return False
    if not any(c in "!$%@#*" for c in password):
        messagebox.showerror("Errore", "La password deve contenere almeno un carattere speciale (! $ % @ # *)")
        return False
    return password

def valida_indirizzo(numero_civico, cap, provincia):
    if len(cap) != 5:
        messagebox.showerror("Errore", "CAP non valido")
        return False
    if len(provincia) != 2:
        messagebox.showerror("Errore", "Provincia non valida")
        return False
    try: 
        numero_civico = int(numero_civico)
        cap = int(cap)
    except ValueError:
        messagebox.showwarning("Attenzione", "Numero civico e CAP devono essere dei valori numerici")
        return False
    return numero_civico, cap, provincia

def valida_prezzo(prezzo_unitario):
    try: 
        prezzo_unitario = float(prezzo_unitario)
    except ValueError:
        messagebox.showwarning("Attenzione", "Prezzo unitario deve essere un valore numerico")
        return None         
    if prezzo_unitario == 0:
        messagebox.showerror("Errore", "Il prezzo unitario non può essere 0")
        return None
    return prezzo_unitario

def valida_quantita_mag(quantita_magazzino):
    try: 
        quantita_magazzino = int(quantita_magazzino)
    except ValueError:
        messagebox.showwarning("Attenzione", "Quantità deve essere un valore numerico")
        return None
    return quantita_magazzino

def valida_quantita_vend(quantita_venduta):
    try:
        quantita_venduta = int(quantita_venduta)
    except ValueError:
        messagebox.showwarning("Attenzione", "La quantità deve essere numerica")
        return None
    if quantita_venduta == 0:
        messagebox.showerror("Errore", "La quantità venduta non può essere 0")
        return None
    return quantita_venduta

def gestione_clienti():
    fin_clienti = tk.Toplevel(root)
    fin_clienti.title("Gestione clienti")
    fin_clienti.geometry(WINDOW_SIZE)
    fin_clienti.config(bg = BG)
    fin_clienti.iconbitmap(FAVICON)

    addWidget(tk.Label(fin_clienti, text = "Gestione clienti", font= (FONT, 16, "bold")), 1, 1, 1, 1)
    addWidget(tk.Label(fin_clienti, text = "ID cliente"), 2, 1, 1, 1)
    entry_id_cliente = ttk.Entry(fin_clienti, width=25)
    addWidget(entry_id_cliente, 2, 2, 1, 1)
    addWidget(tk.Label(fin_clienti, text = "Nome"), 3, 1, 1, 1)
    entry_nome = ttk.Entry(fin_clienti, width=25)
    addWidget(entry_nome, 3, 2, 1, 1)
    addWidget(tk.Label(fin_clienti, text = "Cognome"), 4, 1, 1, 1)
    entry_cognome = ttk.Entry(fin_clienti, width=25)
    addWidget(entry_cognome, 4, 2, 1, 1)
    addWidget(tk.Label(fin_clienti, text = "Email"), 5, 1, 1, 1)
    entry_email = ttk.Entry(fin_clienti, width=25)
    addWidget(entry_email, 5, 2, 1, 1)
    addWidget(tk.Label(fin_clienti, text = "Password"), 6, 1, 1, 1)
    entry_password= ttk.Entry(fin_clienti, show= "*", width=25)
    addWidget(entry_password, 6, 2, 1, 1)
    addWidget(tk.Label(fin_clienti, text = "Indirizzo consegna: "), 7, 1, 1, 1)
    addWidget(tk.Label(fin_clienti, text = "Via"), 8, 1, 1, 1)
    entry_via = ttk.Entry(fin_clienti, width=25)
    addWidget(entry_via, 8, 2, 1, 1)
    addWidget(tk.Label(fin_clienti, text = "Numero civico"), 9, 1, 1, 1)
    entry_numero_civico= ttk.Entry(fin_clienti, width=25)
    addWidget(entry_numero_civico, 9, 2, 1, 1)
    addWidget(tk.Label(fin_clienti, text = "Città"), 10, 1, 1, 1)
    entry_citta= ttk.Entry(fin_clienti, width=25)
    addWidget(entry_citta, 10, 2, 1, 1)
    addWidget(tk.Label(fin_clienti, text = "CAP"), 11, 1, 1, 1)
    entry_cap = ttk.Entry(fin_clienti, width=25)
    addWidget(entry_cap, 11, 2, 1, 1)
    addWidget(tk.Label(fin_clienti, text = "Provincia"), 12, 1, 1, 1)
    entry_provincia= ttk.Entry(fin_clienti, width=25)
    addWidget(entry_provincia, 12, 2, 1, 1)

    tabella = ttk.Treeview(fin_clienti, columns=("id", "nome", "cognome", "email", "indirizzo_consegna"), show="headings", height=12)
    tabella.heading("id", text="ID")
    tabella.heading("nome", text="Nome")
    tabella.heading("cognome", text="Cognome")
    tabella.heading("email", text="Email")
    tabella.heading("indirizzo_consegna", text="Indirizzo")
    tabella.column("id", width= 30)
    tabella.column("nome", width= 100)
    tabella.column("cognome", width= 100)
    tabella.column("email", width= 200)
    tabella.column("indirizzo_consegna", width= 250)
    addWidget(tabella, 2, 4, 12, 6)

    def aggiorna_tabella():
        tabella.delete(*tabella.get_children())
        try:
            cursor.execute("SELECT id_cliente, nome, cognome, email, indirizzo_consegna FROM clienti")
            risultati = cursor.fetchall()
            if not risultati:
                messagebox.showerror("Errore", "Il database è vuoto")
                return
            for ut in risultati:
                tabella.insert("", "end", values=(ut["id_cliente"], ut["nome"], ut["cognome"], ut["email"], ut["indirizzo_consegna"]))
        except mysql.connector.Error:
            messagebox.showerror("Errore", "Errore di connessione con il database")
            return

    def pulire_dati():
        entry_id_cliente.delete(0, "end")
        entry_nome.delete(0, "end")
        entry_cognome.delete(0, "end")
        entry_email.delete(0, "end")
        entry_password.delete(0, "end")
        entry_via.delete(0, "end")
        entry_numero_civico.delete(0, "end")
        entry_citta.delete(0, "end")
        entry_cap.delete(0, "end")
        entry_provincia.delete(0, "end")

    def inserisci_nuovo_cliente():
        id_cliente = entry_id_cliente.get().strip()
        nome = entry_nome.get().strip()
        cognome = entry_cognome.get().strip()
        email = entry_email.get().strip()
        password = entry_password.get().strip()
        via = entry_via.get().strip() 
        numero_civico = entry_numero_civico.get().strip()
        citta = entry_citta.get().strip() 
        cap = entry_cap.get().strip()
        provincia = entry_provincia.get().strip()
        
        if id_cliente:
            messagebox.showerror("Errore", "ID gestito automaticamente dal sistema, lasciare il campo vuoto")
            return
        if not nome or not cognome or not email or not password or not via or not numero_civico or not citta or not cap or not provincia:
            messagebox.showwarning("Attenzione", "Compilare tutti i campi (eccetto il campo ID)")
            return
        email = valida_email(email)
        password = valida_password(password)
        indirizzo = valida_indirizzo(numero_civico, cap, provincia)
        if not email or not password or not indirizzo: 
            return
        hashed = bcrypt.hashpw(password.encode(), bcrypt.gensalt())
        indirizzo_consegna = {"via" : via, "numero_civico" : numero_civico, "citta" : citta, "cap" : cap, "provincia" : provincia}
        indirizzo_consegna_json= json.dumps(indirizzo_consegna)
        nuovo_cliente= Cliente(id_cliente= id_cliente, nome=nome, cognome= cognome, email=email, password= hashed, indirizzo_consegna={"via" : via, "numero_civico" : numero_civico, "citta" : citta, "cap" : cap, "provincia" : provincia})
        try:
            cursor.execute(
                """
                INSERT INTO clienti (nome, cognome, email, password, indirizzo_consegna) 
                VALUES (%s, %s, %s, %s, %s)
                """, (nome, cognome, email, hashed, indirizzo_consegna_json))
            conn.commit()
        except mysql.connector.Error:
            messagebox.showerror("Errore", "Errore di connessione con il database")
            return
        pulire_dati()
        messagebox.showinfo("Info", "Nuovo cliente inserito con successo")
        aggiorna_tabella()

    def modifica_cliente():
        id_cliente = entry_id_cliente.get().strip()
        nome = entry_nome.get().strip()
        cognome = entry_cognome.get().strip()
        email = entry_email.get().strip()
        password = entry_password.get().strip()
        via = entry_via.get().strip() 
        numero_civico = entry_numero_civico.get().strip()
        citta = entry_citta.get().strip() 
        cap = entry_cap.get().strip()
        provincia = entry_provincia.get().strip()   
        if not id_cliente:
            messagebox.showwarning("Attenzione", "Inserire ID")
            return
        cursor.execute("SELECT * FROM clienti WHERE id_cliente = %s", (id_cliente,))
        risultato = cursor.fetchone()
        if not risultato:
            messagebox.showerror("Errore", "Cliente non trovato")
            return
        if nome:
            cursor.execute("""UPDATE clienti
                       SET nome=%s
                       WHERE id_cliente = %s
                       """, (nome, id_cliente))
        if cognome:
            cursor.execute("""UPDATE clienti
                       SET cognome=%s
                       WHERE id_cliente = %s
                       """, (cognome, id_cliente))
        if email:
            email = valida_email(email)
            if not email:
                return
            cursor.execute("""UPDATE clienti
                       SET email=%s
                       WHERE id_cliente = %s
                       """, (email, id_cliente))
        if password:
            password = valida_password(password)
            if not password:
                return
            hashed = bcrypt.hashpw(password.encode(), bcrypt.gensalt())
            cursor.execute("""UPDATE clienti
                       SET password=%s
                       WHERE id_cliente = %s
                       """, (hashed, id_cliente))
        if via and numero_civico and citta and cap and provincia:
            indirizzo = valida_indirizzo(numero_civico, cap, provincia)
            if not indirizzo:
                return
            indirizzo_consegna = {"via": via, "numero_civico": numero_civico, "citta": citta, "cap": cap, "provincia": provincia}
            indirizzo_consegna_json = json.dumps(indirizzo_consegna)
            cursor.execute("""UPDATE clienti
                        SET indirizzo_consegna=%s
                        WHERE id_cliente = %s
                        """, (indirizzo_consegna_json, id_cliente))
        conn.commit()
        messagebox.showinfo("Info", "Dati cliente modificati con successo")
        aggiorna_tabella()
        pulire_dati()

    def elimina_cliente():
        id_cliente = entry_id_cliente.get().strip()
        if not id_cliente:
            messagebox.showwarning("Attenzione", "Inserire ID")
            return
        try:
            if id_cliente:
                cursor.execute("SELECT * FROM clienti WHERE id_cliente = %s", (id_cliente,))
            risultato = cursor.fetchone()
            if not risultato:
                messagebox.showerror("Errore", "Cliente non trovato")
                return
        except mysql.connector.Error:
            messagebox.showerror("Errore", "Errore di connessione con il database")
            return
        cursor.execute("""DELETE FROM clienti
                       WHERE id_cliente = %s
                       """, (id_cliente,))
        conn.commit()
        messagebox.showinfo("Info", "Cliente cancellato")
        pulire_dati()

    def cerca_cliente():
        id_cliente = entry_id_cliente.get().strip()
        if not id_cliente:
            messagebox.showwarning("Attenzione", "Inserire ID")
            return
        try:
            if id_cliente:
                cursor.execute("SELECT * FROM clienti WHERE id_cliente = %s", (id_cliente,))
            risultato = cursor.fetchone()
            if not risultato:
                messagebox.showerror("Errore", "Cliente non trovato")
                return
            tabella.delete(*tabella.get_children())
            tabella.insert("", "end", values=(risultato["id_cliente"], risultato["nome"], risultato["cognome"], risultato["email"], risultato["indirizzo_consegna"]))
        except mysql.connector.Error:
            messagebox.showerror("Errore", "Errore di connessione con il database")
            return
        pulire_dati()

    def visualizza_clienti():
        aggiorna_tabella()

    def esporta_clienti():
        try:
            cursor.execute("SELECT id_cliente, nome, cognome, email, indirizzo_consegna FROM clienti")
            risultati = cursor.fetchall()
            if not risultati:
                messagebox.showerror("Errore", "Database clienti è vuoto")  
                return  
            directory= os.path.join(FILE_PATH / "dati esportati") 
            os.makedirs(directory, exist_ok=True)
            with open(os.path.join(directory, f"Elenco clienti_{date.today()}.csv"), "w", newline="", encoding=ENCODING) as file:
                writer = csv.writer(file)
                writer.writerow(["id cliente", "nome", "cognome", "email", "indirizzo consegna"])
                for ut in risultati:
                    writer.writerow([ ut["id_cliente"], ut["nome"], ut["cognome"], ut["email"], ut["indirizzo_consegna"]])
            messagebox.showinfo("Info", "Il file Elenco clienti è stato creato nella cartella 'dati esportati'")
        except mysql.connector.Error:
            messagebox.showerror("Errore", "Errore di connessione con il database")
            return
    crea_df_clienti()

    addWidget(ttk.Button(fin_clienti, text= "Salva un nuovo cliente", command = inserisci_nuovo_cliente), 16, 1, 1, 1)
    addWidget(ttk.Button(fin_clienti, text= "Modifica dati", command = modifica_cliente), 18, 1, 1, 1)
    addWidget(ttk.Button(fin_clienti, text= "Elimina cliente", command = elimina_cliente), 20, 1, 1, 1)
    addWidget(ttk.Button(fin_clienti, text= "Cerca dati cliente tramite ID", command = cerca_cliente), 16, 2, 1, 1)
    addWidget(ttk.Button(fin_clienti, text= "Visualizza tutti i clienti", command = visualizza_clienti), 18, 2, 1, 1)
    addWidget(ttk.Button(fin_clienti, text= "Esporta elenco clienti", command = esporta_clienti), 20, 2, 1, 1)
    addWidget(ttk.Button(fin_clienti, text= "Chiudi", command = fin_clienti.destroy), 21, 1, 1, 1)


def gestione_prodotti():
    fin_prodotti = tk.Toplevel(root)
    fin_prodotti.title("Gestione prodotti")
    fin_prodotti.geometry(WINDOW_SIZE)
    fin_prodotti.config(bg = BG)
    fin_prodotti.iconbitmap(FAVICON)

    addWidget(tk.Label(fin_prodotti, text = "Gestione prodotti", font= (FONT, 16, "bold")), 1, 1, 1, 1)
    addWidget(tk.Label(fin_prodotti, text = "ID del prodotto"), 2, 1, 1, 1)
    entry_id_prodotto = ttk.Entry(fin_prodotti, width=25)
    addWidget(entry_id_prodotto, 2, 2, 1, 1)
    addWidget(tk.Label(fin_prodotti, text = "Tipo"), 3, 1, 1, 1)
    entry_tipo = ttk.Entry(fin_prodotti, width=25)
    addWidget(entry_tipo, 3, 2, 1, 1)
    addWidget(tk.Label(fin_prodotti, text = "Categoria"), 4, 1, 1, 1)
    entry_categoria = ttk.Entry(fin_prodotti, width=25)
    addWidget(entry_categoria, 4, 2, 1, 1)
    addWidget(tk.Label(fin_prodotti, text = "Descrizione"), 5, 1, 1, 1)
    text_descrizione = tk.Text(fin_prodotti, font=("Calibri", 8), width=25, height= 5)
    addWidget(text_descrizione, 5, 2, 1, 1)
    addWidget(tk.Label(fin_prodotti, text = "Prezzo unitario €"), 6, 1, 1, 1)
    entry_prezzo_unitario = ttk.Entry(fin_prodotti, width=25)
    addWidget(entry_prezzo_unitario, 6, 2, 1, 1)
    addWidget(tk.Label(fin_prodotti, text = "Quantità (n° pezzi)"), 7, 1, 1, 1)
    entry_quantita_magazzino= ttk.Entry(fin_prodotti, width=25)
    addWidget(entry_quantita_magazzino, 7, 2, 1, 1)
    btn_esporta = None 
    tabella = ttk.Treeview(fin_prodotti, columns=("id", "tipo", "categoria", "descrizione", "prezzo_unitario", "quantita_magazzino"), show="headings", height=12)
    tabella.heading("id", text="ID")
    tabella.heading("tipo", text="Tipo")
    tabella.heading("categoria", text="Categoria")
    tabella.heading("descrizione", text="Descrizione")
    tabella.heading("prezzo_unitario", text="Prezzo unitario €")
    tabella.heading("quantita_magazzino", text="Quantità")
    tabella.column("id", width= 30)
    tabella.column("tipo", width= 95)
    tabella.column("categoria", width= 95)
    tabella.column("descrizione", width= 300)
    tabella.column("prezzo_unitario", width= 100)
    tabella.column("quantita_magazzino", width= 80)
    addWidget(tabella, 2, 4, 12, 6)

    def aggiorna_tabella():
        tabella.delete(*tabella.get_children())
        try:
            cursor.execute("SELECT id_prodotto, tipo, categoria, descrizione, prezzo_unitario, quantita_magazzino FROM prodotti")
            risultati = cursor.fetchall()
            for pr in risultati:
                tabella.insert("", "end", values=(pr["id_prodotto"], pr["tipo"], pr["categoria"], pr["descrizione"], pr["prezzo_unitario"], pr["quantita_magazzino"]))
        except mysql.connector.Error:
            messagebox.showerror("Errore", "Errore di connessione con il database")
            return

    def pulire_dati():
        entry_id_prodotto.delete(0, "end")
        entry_tipo.delete(0, "end")
        entry_categoria.delete(0, "end")
        text_descrizione.delete("1.0", "end-1c")
        entry_prezzo_unitario.delete(0, "end")
        entry_quantita_magazzino.delete(0, "end")

    def inserisci_nuovo_prodotto():
        if btn_esporta:
            btn_esporta.destroy()
        id_prodotto = entry_id_prodotto.get().strip()
        tipo = entry_tipo.get().strip()
        categoria = entry_categoria.get().strip()
        descrizione = text_descrizione.get("1.0", "end-1c")
        prezzo_unitario = entry_prezzo_unitario.get().strip()
        quantita_magazzino = entry_quantita_magazzino.get().strip() 

        if id_prodotto:
            messagebox.showerror("Errore", "ID gestito automaticamente dal sistema, lasciare il campo vuoto")
            return
        if not tipo or not categoria or not prezzo_unitario or not quantita_magazzino:
            messagebox.showwarning("Attenzione", "Compilare i campi tipo, categoria, prezzo unitario e quantità")
            return
        prezzo_unitario = valida_prezzo(prezzo_unitario)
        if prezzo_unitario is None:
            return
        quantita_magazzino = valida_quantita_mag(quantita_magazzino)
        if quantita_magazzino is None:
            return
        nuovo_prodotto = Prodotto(id_prodotto= id_prodotto, categoria=categoria, tipo=tipo, descrizione=descrizione, prezzo_unitario=float(prezzo_unitario), quantita_magazzino=int(quantita_magazzino))
        try:
            cursor.execute(
                """
                INSERT INTO prodotti (tipo, categoria, descrizione, prezzo_unitario, quantita_magazzino) 
                VALUES (%s, %s, %s, %s, %s)
                """, (tipo, categoria, descrizione, prezzo_unitario, quantita_magazzino))
            conn.commit()
        except mysql.connector.Error:
            messagebox.showerror("Errore", "Errore di connessione con il database")
            return
        pulire_dati()
        messagebox.showinfo("Info", "Nuovo prodotto inserito con successo")
        aggiorna_tabella()

    def modifica_prodotto():
        if btn_esporta:
            btn_esporta.destroy()
        id_prodotto = entry_id_prodotto.get().strip()
        tipo = entry_tipo.get().strip()
        categoria = entry_categoria.get().strip()
        descrizione = text_descrizione.get("1.0", "end-1c").strip()
        prezzo_unitario = entry_prezzo_unitario.get().strip()
        quantita_magazzino = entry_quantita_magazzino.get().strip() 

        if not id_prodotto:
            messagebox.showwarning("Attenzione", "Inserire ID")
            return
        cursor.execute("SELECT * FROM prodotti WHERE id_prodotto = %s", (id_prodotto,))
        risultato = cursor.fetchone()
        if not risultato:
            messagebox.showerror("Errore", "Prodotto non trovato")
            return
        if tipo:
            cursor.execute("""UPDATE prodotti
                       SET tipo=%s
                       WHERE id_prodotto = %s
                       """, (tipo, id_prodotto))
        if categoria:
            cursor.execute("""UPDATE prodotti
                       SET categoria=%s
                       WHERE id_prodotto = %s
                       """, (categoria, id_prodotto))
        if descrizione:
            cursor.execute("""UPDATE prodotti
                       SET descrizione=%s
                       WHERE id_prodotto = %s
                       """, (descrizione, id_prodotto))
        if prezzo_unitario:
            prezzo_unitario = valida_prezzo(prezzo_unitario)
            if prezzo_unitario is None:
                return
            cursor.execute("""UPDATE prodotti
                       SET prezzo_unitario=%s
                       WHERE id_prodotto = %s
                       """, (prezzo_unitario, id_prodotto))
        if quantita_magazzino:
            quantita_magazzino = valida_quantita_mag(quantita_magazzino)
            if quantita_magazzino is None:
                return    
            cursor.execute("""UPDATE prodotti
                        SET quantita_magazzino=%s
                        WHERE id_prodotto = %s
                        """, (quantita_magazzino, id_prodotto))
        conn.commit()
        messagebox.showinfo("Info", "Dati prodotto modificati con successo")
        aggiorna_tabella()
        pulire_dati()

    def esporta_prodotti():
        if btn_esporta:
            btn_esporta.destroy()
        try:
            cursor.execute("SELECT id_prodotto, tipo, categoria, descrizione, prezzo_unitario, quantita_magazzino FROM prodotti")
            risultati = cursor.fetchall()   
            if not risultati:
                messagebox.showerror("Errore", "Database prodotti è vuoto")  
                return 
            directory= os.path.join(FILE_PATH / "dati esportati") 
            os.makedirs(directory, exist_ok=True) 
            with open(os.path.join(directory, f"Elenco prodotti_{date.today()}.csv"), "w", newline="", encoding=ENCODING) as file:
                writer = csv.writer(file)
                writer.writerow(["id prodotto", "tipo", "categoria", "descrizione", "prezzo unitario", "quantità"])
                for prod in risultati:
                    writer.writerow([prod["id_prodotto"], prod["tipo"], prod["categoria"], prod["descrizione"], prod["prezzo_unitario"], prod["quantita_magazzino"]])
            messagebox.showinfo("Info", "Il file Elenco prodotti è stato creato nella cartella 'dati esportati'")
        except mysql.connector.Error:
            messagebox.showerror("Errore", "Errore di connessione con il database")
            return
    
    def elimina_prodotto():
        if btn_esporta:
            btn_esporta.destroy()
        id_prodotto = entry_id_prodotto.get().strip()
        if not id_prodotto:
            messagebox.showwarning("Attenzione", "Inserire ID")
            return
        try:
            if id_prodotto:
                cursor.execute("SELECT * FROM prodotti WHERE id_prodotto = %s", (id_prodotto,))
            risultato = cursor.fetchone()
            if not risultato:
                messagebox.showerror("Errore", "Prodotto non trovato")
                return
        except mysql.connector.Error:
            messagebox.showerror("Errore", "Errore di connessione con il database")
            return
        cursor.execute("""DELETE FROM prodotti
                       WHERE id_prodotto = %s
                       """, (id_prodotto,))
        conn.commit()
        messagebox.showinfo("Info", "Prodotto cancellato")
        pulire_dati()

    def cerca_prodotto_id():
        if btn_esporta:
            btn_esporta.destroy()
        id_prodotto = entry_id_prodotto.get().strip()
        if not id_prodotto:
            messagebox.showwarning("Attenzione", "Inserire ID")
            return
        try:
            cursor.execute("SELECT * FROM prodotti WHERE id_prodotto = %s", (id_prodotto,))
            risultato = cursor.fetchone()
            if not risultato:
                messagebox.showerror("Errore", "Prodotto non trovato")
                return
            tabella.delete(*tabella.get_children())
            tabella.insert("", "end", values=(risultato["id_prodotto"], risultato["tipo"], risultato["categoria"], risultato["descrizione"], risultato["prezzo_unitario"], risultato["quantita_magazzino"]))
        except mysql.connector.Error:
            messagebox.showerror("Errore", "Errore di connessione con il database")
            return
        pulire_dati()
    
    def cerca_prodotto_categoria():
        nonlocal btn_esporta
        if btn_esporta:
            btn_esporta.destroy()
        categoria= entry_categoria.get().strip()
        if not categoria:
            messagebox.showwarning("Attenzione", "Inserire la categoria")
            return
        try:
            cursor.execute("SELECT * FROM prodotti WHERE categoria=%s", (categoria,))
            risultato = cursor.fetchall()    
            if not risultato:
                messagebox.showerror("Errore", "Prodotti non trovati")
                return
            tabella.delete(*tabella.get_children())
            for prod in risultato:   
                tabella.insert("", "end", values=(prod["id_prodotto"], prod["tipo"], prod["categoria"], prod["descrizione"], prod["prezzo_unitario"], prod["quantita_magazzino"]))
            def esporta():
                directory= os.path.join(FILE_PATH / categoria) 
                os.makedirs(directory, exist_ok=True)
                filename = f"prodotti_categoria_{categoria}_{date.today()}.csv"
                filepath = os.path.join(directory, filename)
                with open(filepath, "w", newline="", encoding=ENCODING) as file:
                    writer = csv.writer(file)
                    writer.writerow(["id prodotto", "tipo", "categoria", "descrizione", "prezzo unitario", "quantità"])
                    for prod in risultato:
                        writer.writerow([prod["id_prodotto"], prod["tipo"], prod["categoria"], prod["descrizione"], prod["prezzo_unitario"], prod["quantita_magazzino"]])
                messagebox.showinfo("Info", f"Il file dei prodotti della categoria {categoria} è stato creato")
            btn_esporta= ttk.Button(fin_prodotti, text="Esporta", command= esporta)
            addWidget(btn_esporta, 18, 9, 1, 1)
        except mysql.connector.Error:
            messagebox.showerror("Errore", "Errore di connessione con il database")
            return
        pulire_dati()

    def visualizza_prodotti():
        if btn_esporta:
            btn_esporta.destroy()
        aggiorna_tabella()
    crea_df_prodotti()

    addWidget(ttk.Button(fin_prodotti, text= "Inserisci un nuovo prodotto", command = inserisci_nuovo_prodotto), 16, 1, 1, 1)
    addWidget(ttk.Button(fin_prodotti, text= "Modifica dati prodotto", command = modifica_prodotto), 18, 1, 1, 1)
    addWidget(ttk.Button(fin_prodotti, text= "Ricerca prodotto per ID", command = cerca_prodotto_id), 16, 2, 1, 1)
    addWidget(ttk.Button(fin_prodotti, text= "Ricerca prodotto per Categoria", command = cerca_prodotto_categoria), 18, 2, 1, 1)
    addWidget(ttk.Button(fin_prodotti, text= "Visualizza tutti i prodotti", command = visualizza_prodotti), 20, 2, 1, 1)
    addWidget(ttk.Button(fin_prodotti, text= "Esporta elenco di tutti i prodotti", command = esporta_prodotti), 20, 1, 1, 1)
    addWidget(ttk.Button(fin_prodotti, text= "Elimina prodotto", command = elimina_prodotto), 21, 1, 1, 1)
    addWidget(ttk.Button(fin_prodotti, text= "Chiudi", command = fin_prodotti.destroy), 22, 1, 1, 1)


def gestione_ordini():
    fin_ordini = tk.Toplevel(root)
    fin_ordini.title("Gestione ordini")
    fin_ordini.geometry(WINDOW_SIZE)
    fin_ordini.config(bg = BG)
    fin_ordini.iconbitmap(FAVICON)

    addWidget(tk.Label(fin_ordini, text = "Gestione ordini", font= (FONT, 16, "bold")), 1, 1, 1, 1)
    addWidget(tk.Label(fin_ordini, text = "ID dell'ordine"), 2, 1, 1, 1)
    entry_id_ordine = ttk.Entry(fin_ordini, width=25)
    addWidget(entry_id_ordine, 2, 2, 1, 1)
    addWidget(tk.Label(fin_ordini, text = "ID del prodotto"), 3, 1, 1, 1)
    entry_id_prodotto= ttk.Entry(fin_ordini, width=25)
    addWidget(entry_id_prodotto, 3, 2, 1, 1)
    addWidget(tk.Label(fin_ordini, text = "ID cliente"), 4, 1, 1, 1)
    entry_id_cliente = ttk.Entry(fin_ordini, width=25)
    addWidget(entry_id_cliente, 4, 2, 1, 1)
    addWidget(tk.Label(fin_ordini, text = "Quantità venduta"), 5, 1, 1, 1)
    entry_quantita_venduta = ttk.Entry(fin_ordini, width=25)
    addWidget(entry_quantita_venduta, 5, 2, 1, 1)
    addWidget(tk.Label(fin_ordini, text = "Data (YYYY-MM-DD)"), 6, 1, 1, 1)
    entry_data_ordine = ttk.Entry(fin_ordini, width=25)
    addWidget(entry_data_ordine, 6, 2, 1, 1)
    
    tabella = ttk.Treeview(fin_ordini, columns=("id_ordine", "id_prodotto", "tipo", "categoria", "id_cliente", "nome", "cognome", "quantita_venduta", "importo", "data_ordine"), show="headings", height=12)
    tabella.heading("id_ordine", text="ID ordine")
    tabella.heading("id_prodotto", text="ID prodotto")
    tabella.heading("tipo", text="Tipo")
    tabella.heading("categoria", text="Categoria")
    tabella.heading("id_cliente", text="ID cliente")
    tabella.heading("nome", text="Nome")
    tabella.heading("cognome", text="Cognome")
    tabella.heading("quantita_venduta", text="Quantità venduta")
    tabella.heading("importo", text="Importo €")
    tabella.heading("data_ordine", text="Data")
    tabella.column("id_ordine", width= 80)
    tabella.column("id_prodotto", width= 80)
    tabella.column("tipo", width= 90)
    tabella.column("categoria", width= 90)
    tabella.column("id_cliente", width= 80)
    tabella.column("nome", width= 90)
    tabella.column("cognome", width= 90)
    tabella.column("quantita_venduta", width= 100)
    tabella.column("importo", width= 90)
    tabella.column("data_ordine", width= 90)
    addWidget(tabella, 2, 4, 12, 6)

    def aggiorna_tabella():
        tabella.delete(*tabella.get_children())
        try:
            cursor.execute("SELECT o.id_ordine, o.id_prodotto, p.tipo, p.categoria, o.id_cliente, c.nome, c.cognome, o.quantita_venduta, o.importo, o.data_ordine FROM ordini o JOIN prodotti p ON o.id_prodotto = p.id_prodotto JOIN clienti c ON o.id_cliente = c.id_cliente")
            risultati = cursor.fetchall()
            if not risultati:
                messagebox.showerror("Errore", "Il database è vuoto")
                return
            for ord in risultati:
                tabella.insert("", "end", values=(ord["id_ordine"], ord["id_prodotto"], ord["tipo"], ord["categoria"], ord["id_cliente"], ord["nome"], ord["cognome"], ord["quantita_venduta"], ord["importo"], ord["data_ordine"]))
        except mysql.connector.Error:
            messagebox.showerror("Errore", "Errore di connessione con il database")
            return

    def pulire_dati():
        entry_id_ordine.delete(0, "end")
        entry_id_prodotto.delete(0, "end")
        entry_id_cliente.delete(0, "end")
        entry_quantita_venduta.delete(0, "end")
        entry_data_ordine.delete(0, "end")

    def inserisci_nuovo_ordine():
        id_ordine = entry_id_ordine.get().strip()
        id_prodotto = entry_id_prodotto.get().strip()
        id_cliente = entry_id_cliente.get().strip()
        quantita_venduta = entry_quantita_venduta.get().strip()
        data_ordine = str(date.today())
        if id_ordine:
            messagebox.showerror("Errore", "ID gestito automaticamente dal sistema, lasciare il campo vuoto")
            return
        if not id_prodotto or not id_cliente or not quantita_venduta:
            messagebox.showwarning("Attenzione", "Compilare i campi ID prodotto, ID cliente e quantità venduta")
            return
        try:
            id_prodotto = int(id_prodotto)
        except ValueError:
            messagebox.showwarning("Attenzione", "ID prodotto deve essere numerico")
            return
        cursor.execute("SELECT quantita_magazzino, prezzo_unitario FROM prodotti WHERE id_prodotto=%s", (id_prodotto,))
        prodotto = cursor.fetchone()
        if not prodotto:
            messagebox.showerror("Errore", "Prodotto non trovato")
            return
        try:
            id_cliente = int(id_cliente)
        except ValueError:
            messagebox.showwarning("Attenzione", "ID cliente deve essere numerico")
            return
        cursor.execute("SELECT * FROM clienti WHERE id_cliente=%s", (id_cliente,))
        if not cursor.fetchone():
            messagebox.showerror("Errore", "Cliente non trovato")
            return
        quantita_venduta = valida_quantita_vend(quantita_venduta)
        quantita_magazzino = prodotto["quantita_magazzino"]
        if quantita_venduta > quantita_magazzino:
            messagebox.showerror("Errore", "Quantità in magazzino insufficiente")
            return
        nuova_quantita = quantita_magazzino - quantita_venduta
        cursor.execute("UPDATE prodotti SET quantita_magazzino=%s WHERE id_prodotto=%s", (nuova_quantita, id_prodotto))
        prezzo_unitario = prodotto["prezzo_unitario"]
        importo = quantita_venduta * prezzo_unitario
        try:
            cursor.execute("""
                INSERT INTO ordini (id_prodotto, id_cliente, quantita_venduta, importo, data_ordine)
                VALUES (%s, %s, %s, %s, %s)
            """, (id_prodotto, id_cliente, quantita_venduta, importo, data_ordine))
            conn.commit()
        except mysql.connector.Error:
            messagebox.showerror("Errore", "Errore di connessione con il database")
            return
        pulire_dati()
        messagebox.showinfo("Info", "Nuovo ordine inserito con successo")
        aggiorna_tabella()

    def modifica_ordine():
        id_ordine = entry_id_ordine.get().strip()
        id_prodotto = entry_id_prodotto.get().strip()
        id_cliente = entry_id_cliente.get().strip()
        quantita_venduta = entry_quantita_venduta.get().strip()
        if not id_ordine:
            messagebox.showwarning("Attenzione", "Inserire l'ID dell'ordine")
            return
        cursor.execute("SELECT * FROM ordini WHERE id_ordine=%s", (id_ordine,))
        ordine = cursor.fetchone()
        if not ordine:
            messagebox.showerror("Errore", "Ordine non trovato")
            return
        id_prodotto_attuale = ordine["id_prodotto"]
        quantita_attuale = ordine["quantita_venduta"]
        if id_prodotto:
            try:
                id_prodotto = int(id_prodotto)
            except ValueError:
                messagebox.showwarning("Attenzione", "L'ID prodotto deve essere numerico")
                return
            cursor.execute("SELECT * FROM prodotti WHERE id_prodotto=%s", (id_prodotto,))
            if not cursor.fetchone():
                messagebox.showerror("Errore", "Prodotto non trovato")
                return
            cursor.execute("UPDATE ordini SET id_prodotto=%s WHERE id_ordine=%s",
                        (id_prodotto, id_ordine))
            id_prodotto_attuale = id_prodotto 
        if id_cliente:
            try:
                id_cliente = int(id_cliente)
            except ValueError:
                messagebox.showwarning("Attenzione", "L'ID cliente deve essere numerico")
                return
            cursor.execute("SELECT * FROM clienti WHERE id_cliente=%s", (id_cliente,))
            if not cursor.fetchone():
                messagebox.showerror("Errore", "Cliente non trovato")
                return
            cursor.execute("UPDATE ordini SET id_cliente=%s WHERE id_ordine=%s", (id_cliente, id_ordine))
        if quantita_venduta:
            quantita_venduta = valida_quantita_vend(quantita_venduta)
            cursor.execute("SELECT quantita_magazzino, prezzo_unitario FROM prodotti WHERE id_prodotto=%s", (id_prodotto_attuale,))
            prodotto = cursor.fetchone()
            if not prodotto:
                messagebox.showerror("Errore", "Prodotto non trovato")
                return
            quantita_magazzino = prodotto["quantita_magazzino"]
            prezzo_unitario = prodotto["prezzo_unitario"]
            quantita_magazzino += quantita_attuale
            if quantita_venduta > quantita_magazzino:
                messagebox.showerror("Errore", "Quantità in magazzino insufficiente")
                return
            nuova_quantita_magazzino = quantita_magazzino - quantita_venduta
            cursor.execute("UPDATE prodotti SET quantita_magazzino=%s WHERE id_prodotto=%s", (nuova_quantita_magazzino, id_prodotto_attuale))
            importo = quantita_venduta * prezzo_unitario
            cursor.execute("""
                UPDATE ordini 
                SET quantita_venduta=%s, importo=%s 
                WHERE id_ordine=%s
            """, (quantita_venduta, importo, id_ordine))
        conn.commit()
        messagebox.showinfo("Info", "Dati ordine modificati con successo")
        aggiorna_tabella()
        pulire_dati()

    def esporta_ordini():
        try:
            cursor.execute("SELECT id_ordine, id_prodotto, id_cliente, quantita_venduta, importo, data_ordine FROM ordini")
            risultati = cursor.fetchall()
            if not risultati:
                messagebox.showerror("Errore", "Database ordini è vuoto")  
                return 
            directory= os.path.join(FILE_PATH / "dati esportati") 
            os.makedirs(directory, exist_ok=True)    
            with open(os.path.join(directory, f"Elenco ordini_{date.today()}.csv"), "w", newline="", encoding=ENCODING) as file:
                writer = csv.writer(file)
                writer.writerow(["id ordine", "id prodotto", "id cliente", "quantita venduta", "importo €", "data dell'ordine"])
                for ord in risultati:
                    writer.writerow([ord["id_ordine"], ord["id_prodotto"], ord["id_cliente"], ord["quantita_venduta"], ord["importo"], ord["data_ordine"]])
            messagebox.showinfo("Info", "Il file Elenco ordini è stato creato nella cartella 'dati esportati'")
        except mysql.connector.Error:
            messagebox.showerror("Errore", "Errore di connessione con il database")
            return
    
    def elimina_ordine():
        id_ordine = entry_id_ordine.get().strip()
        if not id_ordine:
            messagebox.showwarning("Attenzione", "Inserire ID")
            return
        try:
            if id_ordine:
                cursor.execute("SELECT * FROM ordini WHERE id_ordine = %s", (id_ordine,))
                ordine = cursor.fetchone()
            if not ordine:
                messagebox.showerror("Errore", "Ordine non trovato")
                return
        except mysql.connector.Error:
            messagebox.showerror("Errore", "Errore di connessione con il database")
            return
        id_prodotto= ordine['id_prodotto']
        quantita_venduta= ordine['quantita_venduta']
        cursor.execute("SELECT quantita_magazzino FROM prodotti WHERE id_prodotto=%s", (id_prodotto,))
        prodotto = cursor.fetchone()
        if not prodotto:
            messagebox.showerror("Errore", "Prodotto associato all'ordine non trovato")
            return
        quantita_magazzino = prodotto["quantita_magazzino"]
        nuova_quantita = quantita_magazzino + quantita_venduta
        cursor.execute("UPDATE prodotti SET quantita_magazzino=%s WHERE id_prodotto=%s", (nuova_quantita, id_prodotto))
        cursor.execute("DELETE FROM ordini WHERE id_ordine=%s", (id_ordine,))
        conn.commit()
        messagebox.showinfo("Info", "Ordine eliminato con successo")
        pulire_dati()
        aggiorna_tabella()
        
    def cerca_ordine_id():
        id_ordine = entry_id_ordine.get().strip()
        if not id_ordine:
            messagebox.showwarning("Attenzione", "Inserire l'ID dell'ordine")
            return
        try:
            cursor.execute("SELECT * FROM ordini o JOIN prodotti p ON o.id_prodotto = p.id_prodotto JOIN clienti c ON o.id_cliente = c.id_cliente WHERE id_ordine = %s", (id_ordine,))
            risultato = cursor.fetchone()
            if not risultato:
                messagebox.showerror("Errore", "Ordine non trovato")
                return
            tabella.delete(*tabella.get_children())
            tabella.insert("", "end", values=(risultato["id_ordine"], risultato["id_prodotto"], risultato["tipo"], risultato["categoria"], risultato["id_cliente"], risultato["nome"], risultato["cognome"], risultato["quantita_venduta"], risultato["importo"], risultato["data_ordine"]))
        except mysql.connector.Error:
            messagebox.showerror("Errore", "Errore di connessione con il database")
            return
        pulire_dati()

    def cerca_ordine_data():
        data_ordine = entry_data_ordine.get().strip()
        if not data_ordine:
            messagebox.showwarning("Attenzione", "Inserire la data dell'ordine")
            return
        try:
            cursor.execute("SELECT * FROM ordini o JOIN prodotti p ON o.id_prodotto = p.id_prodotto JOIN clienti c ON o.id_cliente = c.id_cliente WHERE data_ordine = %s", (data_ordine,))
            risultati = cursor.fetchall()
            if not risultati:
                messagebox.showerror("Errore", "Ordine non trovato")
                return
            tabella.delete(*tabella.get_children())
            for r in risultati:
                tabella.insert("", "end", values=(r["id_ordine"], r["id_prodotto"], r["tipo"], r["categoria"], r["id_cliente"], r["nome"], r["cognome"], r["quantita_venduta"], r["importo"], r["data_ordine"]))
        except mysql.connector.Error:
            messagebox.showerror("Errore", "Errore di connessione con il database")
            return
        pulire_dati()
    
    def visualizza_ordini():
        aggiorna_tabella()  
    crea_df_ordini()

    addWidget(ttk.Button(fin_ordini, text= "Inserisci un nuovo ordine", command = inserisci_nuovo_ordine), 16, 1, 1, 1)
    addWidget(ttk.Button(fin_ordini, text= "Modifica dati dell'ordine", command = modifica_ordine), 18, 1, 1, 1)
    addWidget(ttk.Button(fin_ordini, text= "Ricerca ordine per ID", command = cerca_ordine_id), 16, 2, 1, 1)
    addWidget(ttk.Button(fin_ordini, text= "Ricerca ordine per data", command = cerca_ordine_data), 18, 2, 1, 1)
    addWidget(ttk.Button(fin_ordini, text= "Visualizza tutti gli ordini", command = visualizza_ordini), 20, 2, 1, 1)
    addWidget(ttk.Button(fin_ordini, text= "Esporta elenco ordini", command = esporta_ordini), 20, 1, 1, 1)
    addWidget(ttk.Button(fin_ordini, text= "Elimina ordine", command = elimina_ordine), 21, 1, 1, 1)
    addWidget(ttk.Button(fin_ordini, text= "Chiudi", command = fin_ordini.destroy), 22, 1, 1, 1)


def analisi_dati():
    fin_analisi = tk.Toplevel(root)
    fin_analisi.title("Analisi dati")
    fin_analisi.geometry(WINDOW_SIZE)
    fin_analisi.config(bg = BG)
    fin_analisi.iconbitmap(FAVICON)

    crea_df()
    global df_completo
    if df_completo is None:
        messagebox.showerror("Errore", "Dati non trovati")
        return
    
    def crea_grafico():
        plt.grid(alpha=0.3)
        plt.xticks(rotation= 45)
        plt.tight_layout()
    
    def crea_report_word(nome_file, titolo, testo, immagini):
        doc = Document()
        doc.add_heading(titolo, level=1)
        for par in testo.split("\n"):
            doc.add_paragraph(par)
        doc.add_page_break()
        doc.add_heading("Grafici", level=2)
        for img in immagini:
            doc.add_picture(img, width=Inches(6))
            doc.add_paragraph("")  
        path= os.path.join(FILE_PATH / "report analisi") 
        directory= os.path.join(path, str(date.today()))
        try:
            os.makedirs(directory, exist_ok=True)
            doc.save(os.path.join(directory, nome_file))
            messagebox.showinfo("Info", "Il report è stato creato nella cartella 'report analisi'")
        except:
            messagebox.showerror("Errore", "Errore durante la creazione del report")
    
    def pulisci_pagina():
        if canvas is not None:
            canvas.destroy() 
    
    def crea_canvas(testo):
        global canvas
        pulisci_pagina() 
        scrollbar = tk.Scrollbar(fin_analisi, orient=tk.VERTICAL)
        addWidget(scrollbar, 3, 23, 8, 1)
        canvas= tk.Canvas(fin_analisi, yscrollcommand=scrollbar.set, width=800, height=600, background="#FFFFFF", highlightbackground="#000000")
        addWidget(canvas, 3, 15, 8, 8)
        scrollbar.config(command=canvas.yview)
        inner_frame = tk.Frame(canvas, bg="#FFFFFF")
        canvas.create_window((0, 0), window=inner_frame, anchor="nw")
        label_testo = tk.Label(inner_frame, text=testo, justify="left", bg="#FFFFFF", font=(FONT, 9), wraplength=600)
        label_testo.pack(anchor="nw", padx=8, pady=8)
        inner_frame.update_idletasks()
        canvas.config(scrollregion=canvas.bbox("all"))

    def report_vendite():
        global df_completo
        df_completo["data_ordine"]= pd.to_datetime(df_completo["data_ordine"])
        df_completo=df_completo.sort_values("data_ordine")
        vendite_totale= df_completo["importo"].sum()
        vendite_giornaliero= df_completo.groupby("data_ordine")["importo"].sum().reset_index()
        vendite_giornaliero= vendite_giornaliero.sort_values("data_ordine")
        vendite_mensile= df_completo.resample('ME', on='data_ordine')["importo"].sum().reset_index()
        vendite_mensile= vendite_mensile.sort_values("data_ordine")
        vendite_categoria= df_completo.groupby("categoria")["importo"].sum().reset_index()
        vendite_medio_totale= df_completo["importo"].mean()
        vendite_mediana_totale= df_completo["importo"].median()
        vendite_std_totale= df_completo["importo"].std()
        vendite_medio_categoriaProd= df_completo.groupby("categoria")["importo"].mean().reset_index()
        vendite_mediana_categoriaProd= df_completo.groupby("categoria")["importo"].median()
        vendite_std_categoriaProd= df_completo.groupby("categoria")["importo"].std()
        stats_fat= df_completo.groupby('categoria')['quantita_venduta'].agg(['mean', 'std']).reset_index()
        path= os.path.join(FILE_PATH / "report analisi") 
        directory= os.path.join(path, str(date.today()))
        os.makedirs(directory, exist_ok=True)
        immagini=[]
        
        if len(vendite_giornaliero) == 0:
            messagebox.showerror("Errore", "Non sono presenti vendite giornaliere")
        else:
            plt.figure(figsize= (8,5))
            plt.plot(vendite_giornaliero["data_ordine"], vendite_giornaliero["importo"], color= "#0008EE", marker="o")
            plt.xlabel("Data ordine")
            plt.ylabel("vendite (€)")
            plt.title("Andamento vendite giornaliero")
            crea_grafico()
            plt.savefig(os.path.join(directory, "Andamento vendite giornaliero.png"))
            plt.close()
            immagini.append(os.path.join(directory, "Andamento vendite giornaliero.png"))
        
        plt.figure(figsize= (8,5))
        plt.bar(vendite_categoria["categoria"], vendite_categoria["importo"], color= "#00A303")
        plt.xlabel("Categorie prodotti")
        plt.ylabel("vendite (€)")
        plt.title("vendite per categorie prodotti")
        crea_grafico()
        plt.savefig(os.path.join(directory, "Vendite per categorie prodotti.png"))
        plt.close()
        immagini.append(os.path.join(directory, "Vendite per categorie prodotti.png"))

        plt.figure(figsize= (8,5))
        plt.bar(stats_fat['categoria'], stats_fat['mean'], yerr= stats_fat['std'], capsize= 5, color='#00A303')
        plt.ylabel("Vendite medie per categoria")
        plt.title("Media e deviazione standard per categoria")
        crea_grafico()
        plt.savefig(os.path.join(directory, "Media e deviazione standard per categoria.png"))
        plt.close()
        immagini.append(os.path.join(directory, "Media e deviazione standard per categoria.png"))
        
        dati = df_completo.groupby("tipo", as_index=False)["importo"].sum()
        dati= dati.sort_values(by= "importo", ascending=False)
        if dati["importo"].sum() == 0:
            messagebox.showerror("Errore", "Impossibile calcolare la curva di lorenz \nImporti pari a 0")
        else:
            dati["quota"]= dati["importo"] / dati["importo"].sum()
            dati["cumulativa"] = dati["quota"].cumsum()
            X= np.linspace(0, 1, len(dati))
            y= dati["cumulativa"].values
            plt.figure(figsize= (8,5))
            plt.plot(np.append([0], X), np.append([0], y), color= "#0008EE", marker="s")
            plt.plot([0, 1], [0, 1], linestyle= "--", color="#F97D02")
            plt.xlabel("Quota dei prodotti")
            plt.ylabel("Quota cumulativa delle vendite")
            plt.title("Grafico di Lorenz delle vendite")
            crea_grafico()
            plt.savefig(os.path.join(directory, "Grafico di Lorenz delle vendite.png"))
            plt.close()
            immagini.append(os.path.join(directory, "Grafico di Lorenz delle vendite.png"))
        if len(dati["importo"]) == 0:
            messagebox.showerror("Errore", "Impossibile calcolare l'indice di Gini \nNon ci sono dati disponibili")
        else:
            def gini(v):
                v= np.array(v)
                v= np.sort(v)
                n= len(v)
                cumulative= np.cumsum(v)
                G= (2 * np.sum((np.arange(1, n+1) * v))) / (n * cumulative[-1]) - (n+1) / n
                return G
            gini_index= gini(dati["importo"])

        plt.figure(figsize= (8,5))
        plt.scatter(df_completo["prezzo_unitario"], df_completo["importo"], color= "#0008EE")
        plt.xlabel("Prezzo")
        plt.ylabel("Vendite")
        plt.title("Grafico di dispersione Prezzo vs. Vendite")
        crea_grafico()
        plt.savefig(os.path.join(directory, "Grafico di dispersione Prezzo vs. Vendite.png"))
        plt.close()
        immagini.append(os.path.join(directory, "Grafico di dispersione Prezzo vs. Vendite.png"))

        correlazione_pearson= df_completo["quantita_venduta"].corr(df_completo["prezzo_unitario"])
        if correlazione_pearson >= 0.9:
            risultato_pearson = "Dipendenza positiva molto forte"
        elif 0.7 <= correlazione_pearson < 0.9:
            risultato_pearson = "Dipendenza positiva forte"
        elif 0.4 <= correlazione_pearson < 0.7:
            risultato_pearson = "Dipendenza positiva moderata"
        elif 0.1 <= correlazione_pearson < 0.4:
            risultato_pearson = "Dipendenza positiva debole"
        elif -0.1 < correlazione_pearson < 0.1:
            risultato_pearson = "Nessuna correlazione lineare"
        elif -0.4 <= correlazione_pearson <= -0.1:
            risultato_pearson = "Dipendenza negativa debole"
        elif -0.7 <= correlazione_pearson < -0.4:
            risultato_pearson = "Dipendenza negativa moderata"
        elif -0.9 <= correlazione_pearson < -0.7:
            risultato_pearson = "Dipendenza negativa forte"
        else:
            risultato_pearson = "Dipendenza negativa molto forte"

        X1= df_completo[['prezzo_unitario']]
        y1= df_completo[['quantita_venduta']]
        modello = LinearRegression()
        modello.fit(X1, y1)
        intercetta= modello.intercept_
        coefficiente= modello.coef_
        for col, coef in zip(X1.columns, modello.coef_):
            risultato_reg_lin=(f"Variabile: {col}, coefficiente: {coef}")

        testo1 = f"""
        vendite totale= {f"{vendite_totale:.2f}"} €\n
        vendite giornaliere= \n{vendite_giornaliero.to_string()}\n
        vendite mensili= \n{vendite_mensile.to_string()}\n
        Media importi= {vendite_medio_totale:,.2f}
        Mediana importi= {vendite_mediana_totale:,.2f}
        Deviazione standard= {vendite_std_totale:,.2f}\n
        Vendite per categoria= \n{vendite_categoria.to_string()} €\n
        Media per categoria= \n{vendite_medio_categoriaProd.to_string()}\n
        Mediana per categoria= \n{vendite_mediana_categoriaProd.to_string()}\n
        Deviazione standard per categoria= \n{vendite_std_categoriaProd.to_string()}\n
        Indice di Gini= {gini_index:,.2f}\n
        Coefficiente di correlazione Pearson {correlazione_pearson:,.2f} = {risultato_pearson}\n
        Risultato del modello di regrezzione lineare= {risultato_reg_lin}\n
        """
        crea_canvas(testo1)

        def genera_report():
            nome_file = f"report_vendite_{date.today()}.docx"
            crea_report_word(nome_file, "Report vendite", testo1, immagini)
        addWidget(ttk.Button(fin_analisi, text="Genera report", command=genera_report), 18, 22, 1, 1)    
        return testo1, immagini
    
    def report_clienti():
        global df_completo
        spesa_totale_cliente= df_completo.groupby("id_cliente")["importo"].sum().reset_index()
        top_clienti= spesa_totale_cliente.sort_values("importo", ascending=False).head(10)
        nuovi_clienti= df_completo.groupby("id_cliente")["data_ordine"].min().reset_index()
        ordini_cliente= df_completo.groupby("id_cliente")["id_ordine"].count()
        nr_medio_ord_cliente= ordini_cliente.mean()
        totale_nr_ordini= df_completo.resample('YE', on='data_ordine')["id_ordine"].count()
        path= os.path.join(FILE_PATH / "report analisi") 
        directory= os.path.join(path, str(date.today()))
        os.makedirs(directory, exist_ok=True)
        immagini=[]
        plt.figure(figsize= (8,5))
        plt.bar(top_clienti["id_cliente"],top_clienti["importo"], color="#5FF662")
        plt.xlabel("Cliente")
        plt.ylabel("Vendite (€)")
        plt.title("Top 10 clienti")
        crea_grafico()
        plt.savefig(os.path.join(directory, "Top 10 clienti.png"))
        plt.close()
        immagini.append(os.path.join(directory, "Top 10 clienti.png"))
        
        testo2= f"""
        Importo speso dai singoli clienti= \n{spesa_totale_cliente.to_string()}\n
        ID dei nuovi clienti= \n{nuovi_clienti.to_string()}\n
        Numero ordini effettuati dai singoli clienti= \n{ordini_cliente.to_string()}\n
        Numero medio di ordini effettuati dai clienti= {nr_medio_ord_cliente}\n
        Totale numero ordini ricevuti= {totale_nr_ordini.to_string()}\n
        """
        crea_canvas(testo2)
        def genera_report():
            nome_file = f"report_clienti_{date.today()}.docx"
            crea_report_word(nome_file, "Report clienti", testo2, immagini)
        addWidget(ttk.Button(fin_analisi, text="Genera report", command=genera_report), 18, 22, 1, 1)    
        return testo2, immagini
    
    def report_prodotti():
        global df_completo
        prodotti_piu_venduti = df_completo.groupby("tipo")["quantita_venduta"].sum().reset_index()
        prodotti_piu_venduti= prodotti_piu_venduti.sort_values("quantita_venduta", ascending= False).head(10)
        prodotti_piu_redditizi= df_completo.groupby("tipo")["importo"].sum().reset_index()
        prodotti_piu_redditizi= prodotti_piu_redditizi.sort_values("importo", ascending= False).head(10)
        prezzo_medio_prodotto= df_completo.groupby("categoria")["prezzo_unitario"].mean()
        quantita_media_prodotto= df_completo.groupby("categoria")["quantita_venduta"].mean()
        totale_quantita_venduta= df_completo["quantita_venduta"].sum()
        media_scorte= df_completo["quantita_magazzino"].mean()
        giacenza_inferiori= df_completo[df_completo["quantita_magazzino"] < 20]
        try:
            rotazione_magazzino = totale_quantita_venduta / media_scorte
        except ZeroDivisionError:
            messagebox.showerror("Errore", "Valori non validi\nImpossibile creare report")
        except ValueError:
            messagebox.showerror("Errore", "Valori non validi\nImpossibile creare report")
            return
        min_vendite_categoriaProd= df_completo.groupby("categoria")["importo"].min()
        max_vendite_categoriaProd= df_completo.groupby("categoria")["importo"].max()
        stats_prod = df_completo.groupby('tipo')['quantita_venduta'].agg(['mean', 'std']).reset_index()
        stats_prod['cv'] = stats_prod['std'] / stats_prod['mean']
        top_prod_stabili = stats_prod.sort_values(by= 'cv').head(10)
        top_prod_volatili =stats_prod.sort_values(by= 'cv', ascending= False).head(10)
        path= os.path.join(FILE_PATH / "report analisi") 
        directory= os.path.join(path, str(date.today())) 
        os.makedirs(directory, exist_ok=True)
        immagini=[]
        plt.figure(figsize= (8,5))
        plt.bar(prodotti_piu_venduti["tipo"],prodotti_piu_venduti["quantita_venduta"], color="#5FF662")
        plt.xlabel("Prodotto")
        plt.ylabel("Quantità venduta")
        plt.title("Top 10 prodotti più venduti")
        crea_grafico()
        plt.savefig(os.path.join(directory, "Top 10 prodotti più venduti.png"))
        plt.close()
        immagini.append(os.path.join(directory, "Top 10 prodotti più venduti.png"))
        
        plt.figure(figsize= (8,5))
        plt.bar(prodotti_piu_redditizi["tipo"],prodotti_piu_redditizi["importo"], color="#5FF662")
        plt.xlabel("Prodotto")
        plt.ylabel("Vendite generato (€)")
        plt.title("Top 10 prodotti più redditizi")
        crea_grafico()
        plt.savefig(os.path.join(directory, "Top 10 prodotti più redditizi.png"))
        plt.close()
        immagini.append(os.path.join(directory, "Top 10 prodotti più redditizi.png"))

        plt.figure(figsize= (8,5))
        plt.bar(prezzo_medio_prodotto.index, prezzo_medio_prodotto.values, color="#00A303")
        plt.title("Prezzo medio per categoria")
        plt.xlabel("Categoria")
        plt.ylabel("Prezzo medio (€)")
        crea_grafico()
        plt.savefig(os.path.join(directory, "Prezzo medio per categoria.png"))
        plt.close()
        immagini.append(os.path.join(directory, "Prezzo medio per categoria.png"))
        
        plt.figure(figsize= (8,5))
        plt.bar(quantita_media_prodotto.index, quantita_media_prodotto.values, color="#00A303")
        plt.title("Quantità media venduta per categoria")
        plt.xlabel("Categoria")
        plt.ylabel("Quantità media")
        crea_grafico()
        plt.savefig(os.path.join(directory, "Quantità media venduta per categoria.png"))
        plt.close()
        immagini.append(os.path.join(directory, "Quantità media venduta per categoria.png"))

        plt.figure(figsize= (8,5))
        plt.bar(top_prod_stabili['tipo'], top_prod_stabili['cv'], color="#5FF662")
        plt.title("Top 10 prodotti più stabili per quantità venduta")
        crea_grafico()
        plt.savefig(os.path.join(directory, "Top 10 prodotti più stabili per quantità venduta.png"))
        plt.close()
        immagini.append(os.path.join(directory, "Top 10 prodotti più stabili per quantità venduta.png"))
    
        plt.figure(figsize= (8,5))
        plt.bar(top_prod_volatili['tipo'], top_prod_volatili['cv'], color="#5FF662")
        plt.title("Top 10 prodotti più volatili per quantità venduta")
        crea_grafico()
        plt.savefig(os.path.join(directory, "Top 10 prodotti più volatili per quantità venduta.png"))
        plt.close()
        immagini.append(os.path.join(directory, "Top 10 prodotti più volatili per quantità venduta.png"))

        plt.figure(figsize= (8,5))
        plt.scatter(stats_prod['mean'], stats_prod['std'], color="#0008EE")
        plt.xlabel("Media quantita venduta")
        plt.ylabel("Deviazione standard")
        plt.title("Media vs. Deviazione standard della quantità vendita")
        crea_grafico()
        plt.savefig(os.path.join(directory, "Media vs. Deviazione standard quantità venduta.png"))
        plt.close()
        immagini.append(os.path.join(directory, "Media vs. Deviazione standard quantità venduta.png"))

        kpi= df_completo.groupby("tipo")["importo"].agg(['mean', 'std']).reset_index()
        kpi = kpi.dropna(subset=['mean', 'std'])
        kpi['cv']= kpi['std']/kpi['mean'] 
        plt.figure(figsize= (8,5))
        plt.scatter(kpi['mean'], kpi['std'], color="#0008EE")
        plt.xlabel("Media vendite")
        plt.ylabel("Deviazione standard")
        plt.title("Media vs. deviazione standard delle vendite")
        crea_grafico()
        plt.savefig(os.path.join(directory, "Media vs. deviazione standard delle vendite.png"))
        immagini.append(os.path.join(directory, "Media vs. deviazione standard delle vendite.png"))
        plt.close()
        X= kpi[['mean', 'std']]
        kmeans = KMeans(n_clusters=3, random_state=30)
        kmeans.fit(X)
        kpi['cluster']= kmeans.labels_
        colors= ['#870000', '#006E02', '#000477']
        plt.figure(figsize= (8,5))
        plt.scatter(kpi['mean'], kpi['std'], c=[colors[c] for c in kpi['cluster']])
        plt.xlabel("Media vendite")
        plt.ylabel("Deviazione standard")
        plt.title("Clustering prodotti e-commerce")
        crea_grafico()
        plt.savefig(os.path.join(directory, "Clustering prodotti e-commerce.png"))
        immagini.append(os.path.join(directory, "Clustering prodotti e-commerce.png"))
        plt.close()
        mapping= {0: "stabile", 1:"medio", 2:"volatile"}
        kpi['classe'] = kpi['cluster'].map(mapping)  
        cv_mean = (kpi['std']/kpi['mean']).mean() 
        if cv_mean > 0.5:
            testo_cv_mean = "Il sistema generale è molto volatile"
        else: 
            testo_cv_mean = "Il sistema generale è stabile"
        top_stabili = kpi.nsmallest(10, 'cv')
        top_volatili = kpi.nlargest(10, 'cv')

        testo3= f"""
        ID Prodotti più venduti= \n{prodotti_piu_venduti.to_string()}\n
        ID Prodotti più redditizi= \n{prodotti_piu_redditizi.to_string()}\n
        Prezzo medio per categoria= \n{prezzo_medio_prodotto.to_string()}\n
        Quantità media venduta per categoria= \n{quantita_media_prodotto.to_string()}\n
        Totale quantità venduta= {totale_quantita_venduta}\n
        Prodotti con giacenza inferiore a 20 pz: \n{giacenza_inferiori}\n
        Rotazione magazzino: {rotazione_magazzino:,.2f}\n
        Minimo vendite per categoria: \n{min_vendite_categoriaProd.to_string()}\n
        Massimo vendite per categoria: \n{max_vendite_categoriaProd.to_string()}\n
        Top prodotti stabili per quantità venduta: \n{top_prod_stabili.to_string()}\n
        Top prodotti volatili per quantità venduta: \n{top_prod_volatili.to_string()}\n
        KPI esaminati:\n
        {kpi[['tipo', 'mean', 'std', 'classe']]}\n
        Coefficiente di variazione medio: {cv_mean:,.2f} = {testo_cv_mean}\n
        Top 10 prodotti stabili in base alle vendite: \n{top_stabili[['tipo', 'mean', 'std', 'classe']]}\n
        Top 10 prodotti volatili in base alle vendite: \n{top_volatili[['tipo', 'mean', 'std', 'classe']]}
        """
        crea_canvas(testo3)
        def genera_report():
            nome_file = f"report_prodotti_{date.today()}.docx"
            crea_report_word(nome_file, "Report prodotti", testo3, immagini)
        addWidget(ttk.Button(fin_analisi, text="Genera report", command=genera_report), 18, 22, 1, 1)    
        return testo3, immagini
    
    def report_completo():
        testo1, img1 = report_vendite()
        testo2, img2 = report_clienti()
        testo3, img3 = report_prodotti()
        testo_completo = testo1 + "\n\n" + testo2 + "\n\n" + testo3
        immagini_compelte= img1 + img2 + img3
        crea_canvas(testo_completo)
        def genera_report():
            nome_file= f"report_completo_{date.today()}.docx"
            crea_report_word(nome_file, "Report completo", testo_completo, immagini_compelte)
        addWidget(ttk.Button(fin_analisi, text="Genera report", command=genera_report), 18, 22, 1, 1)    

    def previsione():
        fin_previsioni=tk.Toplevel(fin_analisi)
        fin_previsioni.title("Previsioni")
        fin_previsioni.geometry(WINDOW_SIZE)
        fin_previsioni.config(bg = BG)
        fin_previsioni.iconbitmap(FAVICON)
        addWidget(tk.Label(fin_previsioni, text="Calcolo previsione delle vendite", font=(FONT, 16, "bold")), 1, 1, 1, 1)
        addWidget(tk.Label(fin_previsioni, text="Metodo base", font="bold"),2, 1, 1, 1)
        addWidget(tk.Label(fin_previsioni, text="Nr. giorni"), 3, 1, 1, 1)
        entry_giorni= ttk.Entry(fin_previsioni, width=25)
        addWidget(entry_giorni, 3, 2, 1, 1)
        addWidget(tk.Label(fin_previsioni, text="Utenti giornalieri"), 4, 1, 1, 1)
        entry_utenti_al_giorno= ttk.Entry(fin_previsioni, width=25)
        addWidget(entry_utenti_al_giorno, 4, 2, 1, 1)
        addWidget(tk.Label(fin_previsioni, text="Nr. vendite previste"), 5, 1, 1, 1)
        entry_profitto_previsto= ttk.Entry(fin_previsioni, width=25)
        addWidget(entry_profitto_previsto, 5, 2, 1, 1) 
        addWidget(tk.Label(fin_previsioni, text="Probabilità di vendita (%)"), 6, 1, 1, 1)
        entry_probabilita= ttk.Entry(fin_previsioni, width=25)
        addWidget(entry_probabilita, 6, 2, 1, 1)
        addWidget(tk.Label(fin_previsioni, text="Metodo Monte Carlo", font="bold"),7, 1, 1, 1)
        addWidget(tk.Label(fin_previsioni, text="Livelli di traffico previsti"),8, 1, 1, 1)
        entry_traffic_level= tk.Entry(fin_previsioni, width=25)
        addWidget(entry_traffic_level, 8, 2, 1, 1)
        addWidget(tk.Label(fin_previsioni, text="Tassi di conversione (%)"), 9, 1, 1, 1)
        entry_conv_rate= tk.Entry(fin_previsioni, width=25)
        addWidget(entry_conv_rate, 9, 2, 1, 1)
        addWidget(tk.Label(fin_previsioni, text="Profitto previsto"), 10, 1, 1, 1)
        entry_profit= tk.Entry(fin_previsioni, width=25)
        addWidget(entry_profit, 10, 2, 1, 1)
        
        def calcola(): 
            try:
                giorni= int(entry_giorni.get())
                utenti_al_giorno= int(entry_utenti_al_giorno.get())
                profitto_previsto= float(entry_profitto_previsto.get())
                probabilita= float(entry_probabilita.get())/100
                traffic_lev= entry_traffic_level.get().split(",")
                conv_rate= entry_conv_rate.get().split(",")
                profit= float(entry_profit.get())
            except ValueError:
                messagebox.showerror("Errore", "Compilare tutti i campi con valori numerici")
                return
            
            risultati=[]
            for _ in range(giorni):
                vendite= 0 
                for _ in range(utenti_al_giorno):
                    if random.random() < probabilita:
                        vendite +=1
                profitto = vendite * profitto_previsto
                risultati.append(profitto)
            try:
                media= round(sum(risultati) / len(risultati), 2)
                minimo= round(min(risultati), 2)
                massimo= round(max(risultati), 2)
            except ZeroDivisionError:
                messagebox.showerror("Errore", "Impossibile calcolare con valori = 0")
                media= "//"
                minimo= "//"
                massimo= "//"
            
            traffic_levels=[]
            conversion_rates=[]
            for tl in traffic_lev:
                try:
                    traffic_levels.append(int(tl))    
                except ValueError:
                    messagebox.showerror("Errore", "Valori livelli di traffico non validi.\nInserire valori numerici separati da ,")
                    return                      
            for cr in conv_rate:
                try:
                    conversion_rates.append(float(cr)/100)
                except ValueError:
                    messagebox.showerror("Errore", "Valori del tasso di conversione non validi.\nInserire valori numerici separati da ,")
                    return
                
            scenarios = list(itertools.product(traffic_levels, conversion_rates))
            results = {}
            for traffic, conv in scenarios:
                revenues = []
                for _ in range(500): 
                    sales = 0
                    for _ in range(traffic):
                        if random.random() < conv:
                            sales += 1 
                    revenues.append(sales * profit)
            try: 
                results[(traffic, conv)] = sum(revenues) / len(revenues)
            except ZeroDivisionError:
                messagebox.showerror("Errore", "Impossibile calcolare il ricavo medio con valori = 0")
            testo_scenari = ""
            for k, v in results.items():
                testo_scenari += f"{k} → {v:,.2f}\n"
            testo=(f"""Previsione base a {giorni} giorni:
Ricavo medio = {media}
Valore minimo = {minimo}
Valore massimo = {massimo}\n
Analisi con metodo "Monte Carlo":
livello di traffico - tasso di conversione - Ricavo medio
{testo_scenari}""")
            canvas= tk.Canvas(fin_previsioni, width=500, height=200, background="#FFFFFF")
            addWidget(canvas, 14, 1, 1, 5)
            canvas.create_text(10, 10, text=testo, anchor="nw", justify="left")
        addWidget(ttk.Button(fin_previsioni, text="Calcola", command=calcola), 12, 5, 1, 1) 
        addWidget(ttk.Button(fin_previsioni, text="Chiudi", command=fin_previsioni.destroy), 13, 5, 1, 1) 

    addWidget(tk.Label(fin_analisi, text = "Analisi dei dati", font= (FONT, 16, "bold")), 1, 1, 1, 1)
    addWidget(ttk.Button(fin_analisi, text= "Esporta tutti i dati", command = esportare_dati), 3, 1, 1, 1)
    addWidget(ttk.Button(fin_analisi, text= "Visualizza dati analisi del vendite", command = report_vendite), 4, 1, 1, 1)
    addWidget(ttk.Button(fin_analisi, text= "Visualizza dati clienti", command = report_clienti), 5, 1, 1, 1)
    addWidget(ttk.Button(fin_analisi, text= "Visualizza dati prodotti", command = report_prodotti), 6, 1, 1, 1)
    addWidget(ttk.Button(fin_analisi, text= "Visualizza report completo", command = report_completo), 7, 1, 1, 1)
    addWidget(ttk.Button(fin_analisi, text= "Calcolo previsioni", command = previsione), 8, 1, 1, 1)
    addWidget(ttk.Button(fin_analisi, text= "Chiudi", command = fin_analisi.destroy), 9, 1, 1, 1)


logo = PhotoImage(file = FILE_PATH / "online-shopping.png")
logo_small = logo.subsample(8, 8)
addWidget(tk.Label(root, image= logo_small), 1, 1, 1, 1)

addWidget(tk.Label(root, text = "Gestionale E-commerce", font = (FONT, 18, "bold")), 1, 2, 1, 1)
addWidget(tk.Label(root, text = "Home", font = (FONT, 16, "bold")), 3, 2, 1, 1)

addWidget(ttk.Button(root, text= "Clienti", command = gestione_clienti), 10, 1, 1, 1)
addWidget(ttk.Button(root, text= "Prodotti", command = gestione_prodotti), 15, 1, 1, 1)
addWidget(ttk.Button(root, text="Ordini", command=gestione_ordini), 20, 1, 1, 1)
addWidget(ttk.Button(root, text="Analisi dei dati", command=analisi_dati), 21, 1, 1, 1)
addWidget(ttk.Button(root, text="Chiudi", command=root.quit), 22, 1, 1, 1)

threading.Thread(target=backup, daemon=True).start()

root.mainloop()
cursor.close()
conn.close()
